"""DocForge service – DOCX template processing and document generation."""

import copy
import io
import re

import mammoth
from docx import Document as DocxDocument


def convert_docx_to_html(docx_bytes: bytes) -> str:
    """Convert a DOCX file to HTML using mammoth.

    Returns an HTML string suitable for rendering in the browser.
    """
    result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
    return result.value


def create_template_docx(
    original_bytes: bytes,
    placeholders: list[dict],
) -> bytes:
    """Create a template DOCX by replacing original text spans with ``{{name}}`` markers.

    Args:
        original_bytes: The raw bytes of the original uploaded DOCX.
        placeholders: List of dicts with at least ``name`` and ``original_text`` keys.

    Returns:
        Bytes of the new template DOCX with placeholder markers inserted.
    """
    doc = DocxDocument(io.BytesIO(original_bytes))

    for ph in placeholders:
        marker = "{{" + ph["name"] + "}}"
        original = ph["original_text"]
        _replace_in_document(doc, original, marker)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_document(
    template_bytes: bytes,
    placeholder_values: dict[str, str],
) -> bytes:
    """Fill a template DOCX by replacing ``{{name}}`` markers with actual values.

    Args:
        template_bytes: Bytes of the template DOCX containing ``{{…}}`` markers.
        placeholder_values: Mapping of placeholder name → value to insert.

    Returns:
        Bytes of the finalised DOCX document.
    """
    doc = DocxDocument(io.BytesIO(template_bytes))

    for name, value in placeholder_values.items():
        marker = "{{" + name + "}}"
        _replace_in_document(doc, marker, value)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_html_preview(
    html_template: str,
    placeholder_values: dict[str, str],
) -> str:
    """Return HTML with ``{{name}}`` markers replaced by provided values.

    Unfilled placeholders are kept as-is so the UI can still highlight them.
    """
    result = html_template
    for name, value in placeholder_values.items():
        result = result.replace("{{" + name + "}}", value)
    return result


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _replace_in_document(doc: DocxDocument, old: str, new: str) -> None:
    """Replace *old* with *new* across paragraphs, tables, headers, and footers."""
    # Body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, old, new)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, old, new)

    # Headers / footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is not None:
                for para in header_footer.paragraphs:
                    _replace_in_paragraph(para, old, new)


def _replace_in_paragraph(paragraph, old: str, new: str) -> None:
    """Replace text across runs in a paragraph, preserving formatting of the first run."""
    full_text = "".join(run.text for run in paragraph.runs)
    if old not in full_text:
        return

    replaced = full_text.replace(old, new)

    # Rebuild runs: put full replaced text in the first run, clear the rest
    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
