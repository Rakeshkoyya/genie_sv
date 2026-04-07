"""Document export service for generating DOCX and TXT files.

Tag System — XML tags used by the LLM response and parsed into Word formatting:

    Tag                                  Renders As
    ───────────────────────────────────  ──────────────────────────────────────────
    <title>Text</title>                  Centered, bold, 20pt — chapter/doc title
    <heading>Text</heading>              Bold, 14pt — section heading
    <subheading>Text</subheading>        Bold, 12pt — subsection heading
    <instruction>Text</instruction>      Italic, 11pt — directions text
    <bold>text</bold>                    Inline bold
    <italic>text</italic>               Inline italic
    <underline>text</underline>          Inline underline
    <label>Q1.</label>                   Bold inline prefix for question numbers
    <blank/> or ___                      Fill-in-the-blank underline (50 chars)
    <hr/>                                Horizontal separator line
    <pagebreak/>                         Word document page break
    <space lines="N"/>                   Vertical writing space (N blank lines)
    <indent>text</indent>                Left-indented paragraph (nestable)
    <box title="Note">content</box>     Bordered box with optional title
    <table>                              Formatted table with bold header row
      <row><cell>A</cell>...</row>
    </table>
    1. Item                              Numbered list
    1.1 Item                             Hierarchical numbered (auto-indented)
    - Item                               Bullet list
"""

import io
import re
from typing import Any
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from xhtml2pdf import pisa


DEFAULT_FONT = "Calibri"
DEFAULT_SIZE = Pt(11)


def sanitize_filename(name: str) -> str:
    """Sanitize a filename by removing invalid characters.
    
    Args:
        name: Original filename
        
    Returns:
        Sanitized filename
    """
    sanitized = re.sub(r"[^a-zA-Z0-9\s\-_]", "", name).strip()
    return sanitized if sanitized else "document"


def _setup_styles(doc: Document) -> None:
    """Configure document default styles."""
    style = doc.styles["Normal"]
    style.font.name = DEFAULT_FONT
    style.font.size = DEFAULT_SIZE
    style.paragraph_format.space_after = Pt(4)


def create_docx(title: str, results: list[dict[str, Any]]) -> bytes:
    """Create a DOCX document from generation results.
    
    Args:
        title: Document title
        results: List of dicts with 'name' and 'content' keys
        
    Returns:
        DOCX file bytes
    """
    doc = Document()
    _setup_styles(doc)
    
    # Add each result (each response already contains its own <title> tag)
    for i, result in enumerate(results):
        content = result.get("content", "")
        _add_parsed_content(doc, content)
        
        # Page break between results
        if i < len(results) - 1:
            doc.add_page_break()
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _normalize_content(text: str) -> str:
    """Pre-process content to normalize tag placement for reliable line-by-line parsing.

    1. Join multi-line ``<indent>`` where the closing tag is on the next line.
    2. Split lines that contain multiple block-level tags into separate lines.
    """
    # 1. Join indent that spans two lines: <indent>content\n  </indent>
    text = re.sub(
        r"(<indent>[^\n]*?)\s*\n\s*</indent>",
        r"\1</indent>",
        text,
    )

    # 2. Split lines with multiple block-level tags onto separate lines
    block_tag = (
        r"(?:<title>.*?</title>)"
        r"|(?:<heading>.*?</heading>)"
        r"|(?:<subheading>.*?</subheading>)"
        r"|(?:<instruction>.*?</instruction>)"
        r"|(?:<indent>.*?</indent>)"
        r"|(?:<hr\s*/>)"
        r"|(?:<pagebreak\s*/>)"
        r"|(?:<blank\s*/>)"
        r"|(?:<space\s+lines=\"\d+\"\s*/>)"
    )
    new_lines: list[str] = []
    for line in text.split("\n"):
        parts = re.split(f"({block_tag})", line.strip())
        segments = [p.strip() for p in parts if p and p.strip()]
        if len(segments) > 1:
            new_lines.extend(segments)
        else:
            new_lines.append(line)
    return "\n".join(new_lines)


def _add_parsed_content(doc: Document, text: str) -> None:
    """Parse response content and add to document with rich formatting.

    Supports XML tags: title, heading, subheading, instruction, bold, italic,
    underline, label, blank, hr, pagebreak, space, indent, box, table.
    Also handles numbered/bullet lists and markdown heading fallback.
    """
    text = _normalize_content(text)
    lines = text.split("\n")
    i = 0

    while i < len(lines):
        trimmed = lines[i].strip()

        # Empty line — minimal gap
        if not trimmed:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run()
            run.font.size = Pt(2)
            i += 1
            continue

        # --- Multi-line blocks ---

        # <table>...</table>
        if trimmed.startswith("<table"):
            block_lines = []
            while i < len(lines) and "</table>" not in lines[i]:
                block_lines.append(lines[i])
                i += 1
            if i < len(lines):
                block_lines.append(lines[i])
                i += 1
            _add_table(doc, "\n".join(block_lines))
            continue

        # <box>...</box>
        if trimmed.startswith("<box"):
            block_lines = []
            while i < len(lines) and "</box>" not in lines[i]:
                block_lines.append(lines[i])
                i += 1
            if i < len(lines):
                block_lines.append(lines[i])
                i += 1
            _add_box(doc, "\n".join(block_lines))
            continue

        # --- Single-line tags ---

        # <title>text</title>
        m = re.match(r"^<title>(.*?)</title>$", trimmed)
        if m:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_inline_formatting(para, m.group(1).strip())
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(20)
                run.font.name = DEFAULT_FONT
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # <heading>text</heading>
        m = re.match(r"^<heading>(.*?)</heading>$", trimmed)
        if m:
            para = doc.add_paragraph()
            _add_inline_formatting(para, m.group(1).strip())
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = DEFAULT_FONT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # <subheading>text</subheading>
        m = re.match(r"^<subheading>(.*?)</subheading>$", trimmed)
        if m:
            para = doc.add_paragraph()
            _add_inline_formatting(para, m.group(1).strip())
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = DEFAULT_FONT
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # <instruction>text</instruction>
        m = re.match(r"^<instruction>(.*?)</instruction>$", trimmed)
        if m:
            para = doc.add_paragraph()
            _add_inline_formatting(para, m.group(1).strip())
            for run in para.runs:
                run.italic = True
                run.font.size = Pt(11)
                run.font.name = DEFAULT_FONT
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # <hr/>
        if trimmed in ("<hr/>", "<hr />"):
            _add_hr(doc)
            i += 1
            continue

        # <pagebreak/>
        if trimmed in ("<pagebreak/>", "<pagebreak />"):
            doc.add_page_break()
            i += 1
            continue

        # <space lines="N"/>
        m = re.match(r'^<space\s+lines="(\d+)"\s*/>$', trimmed)
        if m:
            for _ in range(int(m.group(1))):
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        # <blank/> or ___
        if trimmed in ("<blank/>", "<blank />") or re.match(r"^_{3,}$", trimmed):
            _add_blank_line(doc)
            i += 1
            continue

        # <indent>text</indent> — supports nested indent and inner bullet/number
        # Single-line: <indent>text</indent> or multi-line indent that starts/ends on one line
        m = re.match(r"^((?:<indent>)+)(.*?)((?:</indent>)+)$", trimmed)
        if m:
            open_tags = m.group(1).count("<indent>")
            close_tags = m.group(3).count("</indent>")
            nest_level = min(open_tags, close_tags)
            inner = m.group(2).strip()
            left = Inches(0.4 * nest_level)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = left
            # Add bullet prefix if it's a bullet item
            bm = re.match(r"^[-•*]\s+(.+)", inner)
            if bm:
                bullet_run = para.add_run("•  ")
                bullet_run.font.name = DEFAULT_FONT
                bullet_run.font.size = DEFAULT_SIZE
                _add_inline_formatting(para, bm.group(1))
            else:
                _add_inline_formatting(para, inner)
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Partial <indent> (unclosed on this line) — just indent and strip tag
        if trimmed.startswith("<indent>"):
            nest_level = 0
            inner = trimmed
            while inner.startswith("<indent>"):
                inner = inner[len("<indent>"):]
                nest_level += 1
            # Strip any closing tags that are present
            while inner.endswith("</indent>"):
                inner = inner[:-len("</indent>")]
            inner = inner.strip()
            if inner:
                left = Inches(0.4 * nest_level)
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = left
                bm = re.match(r"^[-•*]\s+(.+)", inner)
                if bm:
                    bullet_run = para.add_run("•  ")
                    bullet_run.font.name = DEFAULT_FONT
                    bullet_run.font.size = DEFAULT_SIZE
                    _add_inline_formatting(para, bm.group(1))
                else:
                    _add_inline_formatting(para, inner)
                para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Stray </indent> from multi-line indent — skip
        if trimmed == "</indent>":
            i += 1
            continue

        # Hierarchical numbered list: 1.1 Item, 1.1.1 Item, etc.
        m = re.match(r"^(\d+(?:\.\d+)+)\s+(.+)", trimmed)
        if m:
            number = m.group(1)
            depth = number.count('.')  # 1.1 = 1, 1.1.1 = 2, etc.
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.4 * depth)
            para.paragraph_format.space_after = Pt(2)
            # Bold number prefix
            num_run = para.add_run(f"{number}  ")
            num_run.bold = True
            num_run.font.name = DEFAULT_FONT
            num_run.font.size = Pt(11) if depth > 0 else DEFAULT_SIZE
            _add_inline_formatting(para, m.group(2))
            i += 1
            continue

        # Numbered list: 1. Item text — rendered as plain paragraph with bold number
        # (avoids Word "List Number" style which has its own indent)
        m = re.match(r"^(\d+)[.)]\s+(.+)", trimmed)
        if m:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0)
            para.paragraph_format.space_after = Pt(2)
            num_run = para.add_run(f"{m.group(1)}. ")
            num_run.bold = True
            num_run.font.name = DEFAULT_FONT
            num_run.font.size = DEFAULT_SIZE
            _add_inline_formatting(para, m.group(2))
            i += 1
            continue

        # Bullet list: - Item text or • Item text
        m = re.match(r"^[-•*]\s+(.+)", trimmed)
        if m:
            para = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(para, m.group(1))
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Markdown headings fallback
        m = re.match(r"^(#{1,3})\s+(.+)", trimmed)
        if m:
            level = len(m.group(1))
            para = doc.add_paragraph()
            run = para.add_run(m.group(2))
            run.bold = True
            run.font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
            run.font.name = DEFAULT_FONT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Regular text with inline formatting
        # Strip any stray block tags that weren't caught
        cleaned = re.sub(r"</?(?:response)[^>]*>", "", trimmed).strip()
        if not cleaned:
            i += 1
            continue
        para = doc.add_paragraph()
        _add_inline_formatting(para, cleaned)
        para.paragraph_format.space_after = Pt(2)
        i += 1


def _add_inline_formatting(para, text: str) -> None:
    """Add text with inline formatting: bold, italic, underline, label, blank."""
    pattern = (
        r"(<bold>[\s\S]*?</bold>"
        r"|<italic>[\s\S]*?</italic>"
        r"|<underline>[\s\S]*?</underline>"
        r"|<label>[\s\S]*?</label>"
        r"|<blank\s*/>"
        r"|\*\*[\s\S]*?\*\*"
        r"|_{3,})"
    )
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        # XML bold
        m = re.match(r"^<bold>([\s\S]*?)</bold>$", part)
        if m:
            run = para.add_run(m.group(1))
            run.bold = True
            run.font.name = DEFAULT_FONT
            continue

        # XML italic
        m = re.match(r"^<italic>([\s\S]*?)</italic>$", part)
        if m:
            run = para.add_run(m.group(1))
            run.italic = True
            run.font.name = DEFAULT_FONT
            continue

        # XML underline
        m = re.match(r"^<underline>([\s\S]*?)</underline>$", part)
        if m:
            run = para.add_run(m.group(1))
            run.underline = True
            run.font.name = DEFAULT_FONT
            continue

        # XML label (bold inline prefix like "Q1.")
        m = re.match(r"^<label>([\s\S]*?)</label>$", part)
        if m:
            run = para.add_run(m.group(1))
            run.bold = True
            run.font.name = DEFAULT_FONT
            continue

        # Inline blank
        if part.strip() in ("<blank/>", "<blank />") or re.match(r"^_{3,}$", part.strip()):
            run = para.add_run("_" * 30)
            run.font.name = DEFAULT_FONT
            continue

        # Markdown bold
        m = re.match(r"^\*\*([\s\S]*?)\*\*$", part)
        if m:
            run = para.add_run(m.group(1))
            run.bold = True
            run.font.name = DEFAULT_FONT
            continue

        # Regular text
        run = para.add_run(part)
        run.font.name = DEFAULT_FONT


def _add_table(doc: Document, block: str) -> None:
    """Parse <table> block and add a formatted table."""
    rows_data = []
    for row_match in re.finditer(r"<row>([\s\S]*?)</row>", block):
        cells = re.findall(r"<cell>([\s\S]*?)</cell>", row_match.group(1))
        if cells:
            rows_data.append(cells)

    if not rows_data:
        return

    max_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=max_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_cells in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row_cells):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            _add_inline_formatting(para, cell_text.strip())
            para.paragraph_format.space_after = Pt(2)

    # Bold the header row
    if len(rows_data) > 1:
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_box(doc: Document, block: str) -> None:
    """Parse <box> block and add bordered content."""
    title_match = re.match(r'<box\s+title="([^"]*)"', block)
    box_title = title_match.group(1) if title_match else None

    inner = re.search(r"<box[^>]*>([\s\S]*?)</box>", block)
    if not inner:
        return

    inner_text = inner.group(1).strip()

    # Split inner content into segments: text chunks and <table> blocks
    segments: list[tuple[str, str]] = []  # ("text", lines) or ("table", block)
    parts = re.split(r"(<table>[\s\S]*?</table>)", inner_text)
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if part.startswith("<table>"):
            segments.append(("table", part))
        else:
            segments.append(("text", part))

    if box_title:
        title_para = doc.add_paragraph()
        run = title_para.add_run(box_title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = DEFAULT_FONT
        _set_paragraph_border(title_para, top=True, left=True, right=True)
        title_para.paragraph_format.space_before = Pt(4)
        title_para.paragraph_format.space_after = Pt(2)
        title_para.paragraph_format.left_indent = Inches(0.15)
        title_para.paragraph_format.right_indent = Inches(0.15)

    # Collect all renderable items to determine first/last for borders
    items: list[tuple[str, Any]] = []  # ("line", text) or ("table", block)
    for seg_type, seg_content in segments:
        if seg_type == "table":
            items.append(("table", seg_content))
        else:
            for ln in seg_content.split("\n"):
                ln = ln.strip()
                if ln:
                    items.append(("line", ln))

    for idx, (item_type, item_content) in enumerate(items):
        is_first = (idx == 0 and not box_title)
        is_last = (idx == len(items) - 1)

        if item_type == "table":
            # Add border paragraph above table if not first
            if not is_first:
                border_para = doc.add_paragraph()
                border_para.paragraph_format.space_before = Pt(0)
                border_para.paragraph_format.space_after = Pt(0)
                _set_paragraph_border(border_para, left=True, right=True)
                border_para.paragraph_format.left_indent = Inches(0.15)
                border_para.paragraph_format.right_indent = Inches(0.15)
            _add_table(doc, item_content)
            # Add border paragraph below table if not last
            if not is_last:
                border_para = doc.add_paragraph()
                border_para.paragraph_format.space_before = Pt(0)
                border_para.paragraph_format.space_after = Pt(0)
                _set_paragraph_border(border_para, left=True, right=True)
                border_para.paragraph_format.left_indent = Inches(0.15)
                border_para.paragraph_format.right_indent = Inches(0.15)
        else:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.15)
            para.paragraph_format.right_indent = Inches(0.15)

            m = re.match(r"^<heading>(.*?)</heading>$", item_content)
            if m:
                _add_inline_formatting(para, m.group(1).strip())
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(14)
                    run.font.name = DEFAULT_FONT
            elif (m := re.match(r"^<subheading>(.*?)</subheading>$", item_content)):
                _add_inline_formatting(para, m.group(1).strip())
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.name = DEFAULT_FONT
            elif (m := re.match(r"^<instruction>(.*?)</instruction>$", item_content)):
                _add_inline_formatting(para, m.group(1).strip())
                for run in para.runs:
                    run.italic = True
                    run.font.name = DEFAULT_FONT
            else:
                _add_inline_formatting(para, item_content)

            _set_paragraph_border(para, top=is_first, bottom=is_last, left=True, right=True)
            para.paragraph_format.space_after = Pt(2) if not is_last else Pt(4)


def _set_paragraph_border(para, top=False, bottom=False, left=False, right=False) -> None:
    """Add borders to a paragraph using OxmlElement."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")

    for side, enabled in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if enabled:
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "4")
            border.set(qn("w:color"), "333333")
            pBdr.append(border)

    pPr.append(pBdr)


def _add_hr(doc: Document) -> None:
    """Add a horizontal rule as a paragraph with a bottom border."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_blank_line(doc: Document) -> None:
    """Add a fill-in-the-blank underline."""
    para = doc.add_paragraph()
    run = para.add_run("_" * 50)
    run.font.name = DEFAULT_FONT
    run.font.size = Pt(11)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)


def create_txt(title: str, content: str) -> bytes:
    """Create a plain text file from content.
    
    Strips XML/HTML tags and normalizes formatting.
    
    Args:
        title: Document title
        content: Content text
        
    Returns:
        TXT file bytes (UTF-8 encoded)
    """
    # Strip XML tags
    text = re.sub(r"</?response>", "", content)
    text = re.sub(r"<title>(.*?)</title>", r"\n\n\1\n", text)
    text = re.sub(r"<heading>(.*?)</heading>", r"\n\n\1\n", text)
    text = re.sub(r"<subheading>(.*?)</subheading>", r"\n\1\n", text)
    text = re.sub(r"<instruction>(.*?)</instruction>", r"\n\1\n", text)
    text = re.sub(r"<bold>(.*?)</bold>", r"\1", text)
    text = re.sub(r"<italic>(.*?)</italic>", r"\1", text)
    text = re.sub(r"<underline>(.*?)</underline>", r"\1", text)
    text = re.sub(r"<label>(.*?)</label>", r"\1", text)
    text = re.sub(r"<indent>(.*?)</indent>", r"  \1", text)
    text = re.sub(r"</?box[^>]*>", "", text)
    text = re.sub(r"</?table>", "", text)
    text = re.sub(r"</?row>", "", text)
    text = re.sub(r"<cell>(.*?)</cell>", r"\1\t", text)
    text = re.sub(r"<hr\s*/>", "\n---\n", text)
    text = re.sub(r"<pagebreak\s*/>", "\n\n", text)
    text = re.sub(r"<blank\s*/>", "_______________", text)
    text = re.sub(r'<space\s+lines="\d+"\s*/>', "\n\n", text)
    
    # Build final text
    final_text = f"{title}\n{'=' * len(title)}\n\n{text.strip()}"
    
    return final_text.encode("utf-8")


# ---------------------------------------------------------------------------
# PDF generation — converts XML tags to styled HTML, then renders via xhtml2pdf
# ---------------------------------------------------------------------------

_PDF_CSS = """
@page {
    size: A4;
    margin: 2cm;
}
body {
    font-family: Helvetica, Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.5;
    color: #1a1a1a;
}
h1.doc-title {
    font-size: 20pt;
    font-weight: bold;
    text-align: center;
    margin: 12pt 0 8pt;
}
h2.doc-heading {
    font-size: 14pt;
    font-weight: bold;
    margin: 10pt 0 4pt;
}
h3.doc-subheading {
    font-size: 12pt;
    font-weight: bold;
    margin: 8pt 0 4pt;
}
p.doc-instruction {
    font-style: italic;
    color: #555;
    margin: 4pt 0;
}
p {
    margin: 3pt 0;
}
hr {
    border: none;
    border-top: 1px solid #999;
    margin: 8pt 0;
}
.blank-line {
    letter-spacing: 2pt;
    color: #666;
}
.box {
    border: 1px solid #333;
    border-radius: 4pt;
    padding: 8pt 10pt;
    margin: 8pt 0;
    background-color: #fafafa;
}
.box-title {
    font-weight: bold;
    font-size: 11pt;
    margin-bottom: 4pt;
}
table {
    width: 100%;
    border-collapse: collapse;
    margin: 8pt 0;
}
td, th {
    border: 1px solid #ccc;
    padding: 4pt 8pt;
    text-align: left;
    font-size: 10pt;
}
th {
    background-color: #f0f0f0;
    font-weight: bold;
}
ul, ol {
    margin: 4pt 0 4pt 18pt;
    padding: 0;
}
li {
    margin: 2pt 0;
}
.indent-1 { margin-left: 24pt; }
.indent-2 { margin-left: 48pt; }
.indent-3 { margin-left: 72pt; }
.hier-num { font-weight: bold; margin-right: 4pt; }
.page-break { page-break-before: always; }
"""


def create_pdf(title: str, results: list[dict[str, Any]]) -> bytes:
    """Create a PDF document from generation results.

    Converts XML-tagged content to styled HTML, then renders to PDF
    via xhtml2pdf — producing output that closely matches the preview.

    Args:
        title: Document title
        results: List of dicts with 'name' and 'content' keys

    Returns:
        PDF file bytes
    """
    body_parts: list[str] = []
    for i, result in enumerate(results):
        content = result.get("content", "")
        body_parts.append(_content_to_html(content))
        if i < len(results) - 1:
            body_parts.append('<div class="page-break"></div>')

    html = (
        "<!DOCTYPE html><html><head>"
        "<meta charset='utf-8'/>"
        f"<style>{_PDF_CSS}</style>"
        "</head><body>"
        + "\n".join(body_parts)
        + "</body></html>"
    )

    buffer = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html), dest=buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _content_to_html(text: str) -> str:
    """Convert XML-tagged content to HTML matching the preview styling."""
    text = _normalize_content(text)
    lines = text.split("\n")
    parts: list[str] = []
    i = 0

    while i < len(lines):
        trimmed = lines[i].strip()

        # Empty line
        if not trimmed:
            parts.append("<br/>")
            i += 1
            continue

        # <table>...</table>
        if trimmed.startswith("<table"):
            block_lines: list[str] = []
            while i < len(lines) and "</table>" not in lines[i]:
                block_lines.append(lines[i])
                i += 1
            if i < len(lines):
                block_lines.append(lines[i])
                i += 1
            parts.append(_table_to_html("\n".join(block_lines)))
            continue

        # <box>...</box>
        if trimmed.startswith("<box"):
            block_lines = []
            while i < len(lines) and "</box>" not in lines[i]:
                block_lines.append(lines[i])
                i += 1
            if i < len(lines):
                block_lines.append(lines[i])
                i += 1
            parts.append(_box_to_html("\n".join(block_lines)))
            continue

        # <title>
        m = re.match(r"^<title>(.*?)</title>$", trimmed)
        if m:
            parts.append(f'<h1 class="doc-title">{_inline_to_html(m.group(1).strip())}</h1>')
            i += 1
            continue

        # <heading>
        m = re.match(r"^<heading>(.*?)</heading>$", trimmed)
        if m:
            parts.append(f'<h2 class="doc-heading">{_inline_to_html(m.group(1).strip())}</h2>')
            i += 1
            continue

        # <subheading>
        m = re.match(r"^<subheading>(.*?)</subheading>$", trimmed)
        if m:
            parts.append(f'<h3 class="doc-subheading">{_inline_to_html(m.group(1).strip())}</h3>')
            i += 1
            continue

        # <instruction>
        m = re.match(r"^<instruction>(.*?)</instruction>$", trimmed)
        if m:
            parts.append(f'<p class="doc-instruction">{_inline_to_html(m.group(1).strip())}</p>')
            i += 1
            continue

        # <hr/>
        if trimmed in ("<hr/>", "<hr />"):
            parts.append("<hr/>")
            i += 1
            continue

        # <pagebreak/>
        if trimmed in ("<pagebreak/>", "<pagebreak />"):
            parts.append('<div class="page-break"></div>')
            i += 1
            continue

        # <space lines="N"/>
        m = re.match(r'^<space\s+lines="(\d+)"\s*/>$', trimmed)
        if m:
            for _ in range(int(m.group(1))):
                parts.append("<br/>")
            i += 1
            continue

        # <blank/>
        if trimmed in ("<blank/>", "<blank />") or re.match(r"^_{3,}$", trimmed):
            parts.append(f'<p class="blank-line">{"_" * 40}</p>')
            i += 1
            continue

        # <indent>...</indent>
        if trimmed.startswith("<indent>"):
            inner = trimmed
            level = 0
            while inner.startswith("<indent>"):
                inner = inner[len("<indent>"):]
                level += 1
            while inner.endswith("</indent>"):
                inner = inner[:-len("</indent>")]
            inner = inner.strip()
            if inner:
                cls = f"indent-{min(level, 3)}"
                bm = re.match(r"^[-•*]\s+(.+)", inner)
                if bm:
                    parts.append(f'<p class="{cls}">• {_inline_to_html(bm.group(1))}</p>')
                else:
                    parts.append(f'<p class="{cls}">{_inline_to_html(inner)}</p>')
            i += 1
            continue

        # Stray </indent> from multi-line indent — skip
        if trimmed == "</indent>":
            i += 1
            continue

        # Hierarchical numbered: 1.1, 1.1.1
        m = re.match(r"^(\d+(?:\.\d+)+)\s+(.+)", trimmed)
        if m:
            depth = m.group(1).count(".")
            cls = f"indent-{min(depth, 3)}"
            parts.append(
                f'<p class="{cls}"><span class="hier-num">{m.group(1)}</span> '
                f"{_inline_to_html(m.group(2))}</p>"
            )
            i += 1
            continue

        # Numbered list: 1. Item
        m = re.match(r"^(\d+)[.)]\s+(.+)", trimmed)
        if m:
            parts.append(
                f"<p><strong>{m.group(1)}.</strong> {_inline_to_html(m.group(2))}</p>"
            )
            i += 1
            continue

        # Bullet list
        m = re.match(r"^[-•*]\s+(.+)", trimmed)
        if m:
            parts.append(f"<p style='margin-left:18pt;'>• {_inline_to_html(m.group(1))}</p>")
            i += 1
            continue

        # Markdown headings fallback
        m = re.match(r"^(#{1,3})\s+(.+)", trimmed)
        if m:
            lvl = len(m.group(1))
            tag = f"h{lvl + 1}"
            parts.append(f"<{tag}>{_inline_to_html(m.group(2))}</{tag}>")
            i += 1
            continue

        # Regular text
        cleaned = re.sub(r"</?(?:response)[^>]*>", "", trimmed).strip()
        if cleaned:
            parts.append(f"<p>{_inline_to_html(cleaned)}</p>")
        i += 1

    return "\n".join(parts)


def _inline_to_html(text: str) -> str:
    """Convert inline tags to HTML."""
    # Escape HTML entities first (but preserve our tags)
    # We process tags by splitting, so escape only the plain-text parts
    pattern = (
        r"(<bold>[\s\S]*?</bold>"
        r"|<italic>[\s\S]*?</italic>"
        r"|<underline>[\s\S]*?</underline>"
        r"|<label>[\s\S]*?</label>"
        r"|<blank\s*/>"
        r"|\*\*[\s\S]*?\*\*"
        r"|_{3,})"
    )
    parts = re.split(pattern, text)
    out: list[str] = []

    for part in parts:
        if not part:
            continue

        m = re.match(r"^<bold>([\s\S]*?)</bold>$", part)
        if m:
            out.append(f"<strong>{_esc(m.group(1))}</strong>")
            continue

        m = re.match(r"^<italic>([\s\S]*?)</italic>$", part)
        if m:
            out.append(f"<em>{_esc(m.group(1))}</em>")
            continue

        m = re.match(r"^<underline>([\s\S]*?)</underline>$", part)
        if m:
            out.append(f"<u>{_esc(m.group(1))}</u>")
            continue

        m = re.match(r"^<label>([\s\S]*?)</label>$", part)
        if m:
            out.append(f"<strong>{_esc(m.group(1))}</strong>")
            continue

        if part.strip() in ("<blank/>", "<blank />") or re.match(r"^_{3,}$", part.strip()):
            out.append(f'<span class="blank-line">{"_" * 30}</span>')
            continue

        m = re.match(r"^\*\*([\s\S]*?)\*\*$", part)
        if m:
            out.append(f"<strong>{_esc(m.group(1))}</strong>")
            continue

        out.append(_esc(part))

    return "".join(out)


def _esc(text: str) -> str:
    """Escape HTML special characters."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _table_to_html(block: str) -> str:
    """Convert <table> block to HTML table."""
    rows_data: list[list[str]] = []
    for row_match in re.finditer(r"<row>([\s\S]*?)</row>", block):
        cells = re.findall(r"<cell>([\s\S]*?)</cell>", row_match.group(1))
        if cells:
            rows_data.append(cells)

    if not rows_data:
        return ""

    html_parts = ["<table>"]
    for r_idx, row_cells in enumerate(rows_data):
        html_parts.append("<tr>")
        tag = "th" if r_idx == 0 and len(rows_data) > 1 else "td"
        for cell_text in row_cells:
            html_parts.append(f"<{tag}>{_inline_to_html(cell_text.strip())}</{tag}>")
        html_parts.append("</tr>")
    html_parts.append("</table>")
    return "\n".join(html_parts)


def _box_to_html(block: str) -> str:
    """Convert <box> block to HTML."""
    title_match = re.match(r'<box\s+title="([^"]*)"', block)
    inner = re.search(r"<box[^>]*>([\s\S]*?)</box>", block)
    if not inner:
        return ""

    inner_text = inner.group(1).strip()
    parts = ['<div class="box">']
    if title_match:
        parts.append(f'<div class="box-title">{_esc(title_match.group(1))}</div>')

    # Handle inner tables
    segments = re.split(r"(<table>[\s\S]*?</table>)", inner_text)
    for seg in segments:
        seg = seg.strip()
        if not seg:
            continue
        if seg.startswith("<table>"):
            parts.append(_table_to_html(seg))
        else:
            for ln in seg.split("\n"):
                ln = ln.strip()
                if not ln:
                    continue
                m = re.match(r"^<heading>(.*?)</heading>$", ln)
                if m:
                    parts.append(f'<h2 class="doc-heading">{_inline_to_html(m.group(1).strip())}</h2>')
                elif (m := re.match(r"^<subheading>(.*?)</subheading>$", ln)):
                    parts.append(f'<h3 class="doc-subheading">{_inline_to_html(m.group(1).strip())}</h3>')
                elif (m := re.match(r"^<instruction>(.*?)</instruction>$", ln)):
                    parts.append(f'<p class="doc-instruction">{_inline_to_html(m.group(1).strip())}</p>')
                elif ln in ("<hr/>", "<hr />"):
                    parts.append("<hr/>")
                else:
                    parts.append(f"<p>{_inline_to_html(ln)}</p>")

    parts.append("</div>")
    return "\n".join(parts)


class DocumentExportService:
    """Service class for document export operations."""
    
    def create_docx(self, title: str, results: list[dict[str, Any]]) -> bytes:
        """Create a DOCX document.
        
        Args:
            title: Document title
            results: List of dicts with 'name' and 'content' keys
            
        Returns:
            DOCX file bytes
        """
        return create_docx(title, results)
    
    def create_txt(self, title: str, content: str) -> bytes:
        """Create a plain text document.
        
        Args:
            title: Document title
            content: Content text
            
        Returns:
            TXT file bytes
        """
        return create_txt(title, content)

    def create_pdf(self, title: str, results: list[dict[str, Any]]) -> bytes:
        """Create a PDF document.
        
        Args:
            title: Document title
            results: List of dicts with 'name' and 'content' keys
            
        Returns:
            PDF file bytes
        """
        return create_pdf(title, results)
