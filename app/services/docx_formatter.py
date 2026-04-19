"""DocAgent Word document formatter layer.

Converts the structured DocumentSchema JSON into a professional .docx file
using python-docx with full formatting support:
  - Colored headings at multiple levels
  - Rich text paragraphs (bold, italic, underline, color, highlight)
  - Tables with colored headers, striped rows, borders
  - Bullet and numbered lists
  - Page breaks, horizontal rules, spacers
"""

import io
import logging
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from app.schemas.docagent import DocumentSchema, DocumentBlock, TextRun, TableCell

logger = logging.getLogger(__name__)


def _hex_to_rgb(hex_color: str | None) -> RGBColor | None:
    """Convert #RRGGBB hex string to RGBColor."""
    if not hex_color:
        return None
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        return None
    try:
        r, g, b = int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
    except ValueError:
        return None


def _align(value: str | None) -> WD_ALIGN_PARAGRAPH | None:
    """Convert alignment string to WD_ALIGN_PARAGRAPH."""
    if not value:
        return None
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(value.lower())


def _set_cell_shading(cell, hex_color: str):
    """Set background shading on a table cell."""
    color = hex_color.lstrip("#")
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _apply_run_format(run, text_run: TextRun | dict):
    """Apply formatting from a TextRun to a docx run."""
    if isinstance(text_run, dict):
        text_run = TextRun(**text_run)

    run.bold = text_run.bold
    run.italic = text_run.italic
    run.underline = text_run.underline

    color = _hex_to_rgb(text_run.color)
    if color:
        run.font.color.rgb = color

    if text_run.font_size:
        run.font.size = Pt(text_run.font_size)

    if text_run.highlight:
        # python-docx doesn't have direct highlight API — use character shading
        hex_val = text_run.highlight.lstrip("#")
        rpr = run._r.get_or_add_rPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_val}" w:val="clear"/>')
        rpr.append(shading)


def _add_horizontal_rule(doc: Document):
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
        "</w:pBdr>"
    )
    pPr.append(pBdr)


def _set_table_borders(table, color: str = "CCCCCC"):
    """Set borders on a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f"<w:tblPr {nsdecls('w')}/>")

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f"</w:tblBorders>"
    )
    tblPr.append(borders)


def _normalize_cell(cell_data) -> TableCell:
    """Normalize a table cell — could be a dict, TableCell, or plain string."""
    if isinstance(cell_data, TableCell):
        return cell_data
    if isinstance(cell_data, dict):
        return TableCell(**cell_data)
    return TableCell(text=str(cell_data))


class DocxFormatter:
    """Converts a DocumentSchema into a .docx file."""

    def format(self, schema: DocumentSchema) -> bytes:
        """Generate a .docx file from the document schema.

        Returns:
            The .docx file as bytes.
        """
        doc = Document()

        # ── Page margins ──
        for section in doc.sections:
            margin = Inches(schema.page_margin_inches)
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        # ── Default font — compact body text ──
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(9)
        font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(1)
        style.paragraph_format.line_spacing = Pt(12)

        # ── Configure heading styles — smaller and tighter ──
        for level in range(1, 4):
            heading_style = doc.styles[f"Heading {level}"]
            heading_style.font.name = "Calibri"
            heading_style.font.bold = True
            sizes = {1: 12, 2: 10, 3: 9}
            heading_style.font.size = Pt(sizes.get(level, 10))
            heading_style.paragraph_format.space_before = Pt(6 if level == 1 else 4)
            heading_style.paragraph_format.space_after = Pt(1)
            heading_style.paragraph_format.line_spacing = Pt(14 if level == 1 else 12)

        # ── List styles — tighter ──
        for style_name in ("List Bullet", "List Number"):
            try:
                lst_style = doc.styles[style_name]
                lst_style.font.name = "Calibri"
                lst_style.font.size = Pt(9)
                lst_style.paragraph_format.space_before = Pt(0)
                lst_style.paragraph_format.space_after = Pt(1)
                lst_style.paragraph_format.line_spacing = Pt(12)
            except KeyError:
                pass

        # ── Title ──
        if schema.title:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_para.paragraph_format.space_after = Pt(1)
            title_para.paragraph_format.space_before = Pt(0)
            run = title_para.add_run(schema.title.text)
            run.bold = True
            run.font.size = Pt(min(schema.title.font_size or 14, 14))
            run.font.name = "Calibri"
            color = _hex_to_rgb(schema.title.color)
            if color:
                run.font.color.rgb = color

        # ── Subtitle ──
        if schema.subtitle:
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sub_para.paragraph_format.space_after = Pt(4)
            sub_para.paragraph_format.space_before = Pt(0)
            run = sub_para.add_run(schema.subtitle.text)
            run.italic = True
            run.font.size = Pt(min(schema.subtitle.font_size or 9, 10))
            run.font.name = "Calibri"
            color = _hex_to_rgb(schema.subtitle.color)
            if color:
                run.font.color.rgb = color

        # ── Blocks ──
        for block in schema.blocks:
            self._render_block(doc, block)

        # ── Serialize ──
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ─────────────────────────────────────────────────────────────
    #  Block renderers
    # ─────────────────────────────────────────────────────────────

    def _render_block(self, doc: Document, block: DocumentBlock):
        block_type = block.type.lower().replace(" ", "_")
        renderer = {
            "heading": self._render_heading,
            "paragraph": self._render_paragraph,
            "table": self._render_table,
            "bullet_list": self._render_bullet_list,
            "numbered_list": self._render_numbered_list,
            "page_break": self._render_page_break,
            "horizontal_rule": self._render_horizontal_rule,
            "spacer": self._render_spacer,
        }.get(block_type)

        if renderer:
            renderer(doc, block)
        else:
            logger.warning("Unknown block type: %s", block.type)

    def _render_heading(self, doc: Document, block: DocumentBlock):
        level = min(block.level or 1, 3)
        heading = doc.add_heading(level=level)
        alignment = _align(block.alignment)
        if alignment is not None:
            heading.alignment = alignment
        run = heading.add_run(block.text or "")
        run.bold = True
        if block.underline:
            run.underline = True
        color = _hex_to_rgb(block.color)
        if color:
            run.font.color.rgb = color

    def _render_paragraph(self, doc: Document, block: DocumentBlock):
        para = doc.add_paragraph()
        alignment = _align(block.alignment)
        if alignment is not None:
            para.alignment = alignment
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.line_spacing = Pt(12)

        for run_data in (block.runs or []):
            run = para.add_run(run_data.text if isinstance(run_data, TextRun) else run_data.get("text", ""))
            _apply_run_format(run, run_data)

    def _render_table(self, doc: Document, block: DocumentBlock):
        headers = [_normalize_cell(h) for h in (block.headers or [])]
        rows = [[_normalize_cell(c) for c in row] for row in (block.rows or [])]

        if not headers and not rows:
            return

        col_count = max(len(headers), max((len(r) for r in rows), default=0))
        row_count = (1 if headers else 0) + len(rows)

        table = doc.add_table(rows=row_count, cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True

        # Compact cell margins via table-level XML
        tbl_pr = table._tbl.tblPr if table._tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        cell_mar = parse_xml(
            f'<w:tblCellMar {nsdecls("w")}>'
            '  <w:top w:w="20" w:type="dxa"/>'
            '  <w:left w:w="40" w:type="dxa"/>'
            '  <w:bottom w:w="20" w:type="dxa"/>'
            '  <w:right w:w="40" w:type="dxa"/>'
            '</w:tblCellMar>'
        )
        tbl_pr.append(cell_mar)

        border_hex = (block.border_color or "#CCCCCC").lstrip("#")
        _set_table_borders(table, border_hex)

        # Header row
        if headers:
            for idx, cell_data in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.text = ""
                para = cell.paragraphs[0]
                para.alignment = _align(cell_data.alignment) or WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(cell_data.text)
                run.bold = True
                run.font.size = Pt(8)
                run.font.name = "Calibri"
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                hdr_color = _hex_to_rgb(cell_data.color or "#FFFFFF")
                if hdr_color:
                    run.font.color.rgb = hdr_color
                bg = cell_data.bg_color or block.header_bg_color or "#2C3E50"
                _set_cell_shading(cell, bg)

        # Data rows
        start_row = 1 if headers else 0
        stripe_color = "#F8F9FA"
        for row_idx, row_data in enumerate(rows):
            table_row = table.rows[start_row + row_idx]
            for col_idx, cell_data in enumerate(row_data):
                if col_idx >= col_count:
                    break
                cell = table_row.cells[col_idx]
                cell.text = ""
                para = cell.paragraphs[0]
                para.alignment = _align(cell_data.alignment) or WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(cell_data.text)
                run.font.size = Pt(8)
                run.font.name = "Calibri"
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                if cell_data.bold:
                    run.bold = True
                color = _hex_to_rgb(cell_data.color)
                if color:
                    run.font.color.rgb = color
                if cell_data.bg_color:
                    _set_cell_shading(cell, cell_data.bg_color)
                elif block.striped and row_idx % 2 == 1:
                    _set_cell_shading(cell, stripe_color)

        # Minimal spacing after table
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(1)

    def _render_bullet_list(self, doc: Document, block: DocumentBlock):
        for item in (block.items or []):
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.line_spacing = Pt(12)
            item_data = item if isinstance(item, TextRun) else TextRun(**(item if isinstance(item, dict) else {"text": str(item)}))
            run = para.add_run(item_data.text)
            _apply_run_format(run, item_data)

    def _render_numbered_list(self, doc: Document, block: DocumentBlock):
        for item in (block.items or []):
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.line_spacing = Pt(12)
            item_data = item if isinstance(item, TextRun) else TextRun(**(item if isinstance(item, dict) else {"text": str(item)}))
            run = para.add_run(item_data.text)
            _apply_run_format(run, item_data)

    def _render_page_break(self, doc: Document, block: DocumentBlock):
        doc.add_page_break()

    def _render_horizontal_rule(self, doc: Document, block: DocumentBlock):
        _add_horizontal_rule(doc)

    def _render_spacer(self, doc: Document, block: DocumentBlock):
        # In compact mode, spacers are minimal — 1 tiny gap regardless of requested lines
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
